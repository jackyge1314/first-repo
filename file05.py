from docx import Document
from docx.table import Table

def extract_all_content(file_path):
    doc = Document(file_path)
    content = []
    
    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
    
    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    content.append(cell.text)
    
    return "\n".join(content) if content else "文件无文本内容"

text = extract_all_content("example/test.doc")
print("提取结果：")
print(text)